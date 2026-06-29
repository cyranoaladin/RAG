
from __future__ import annotations

import hashlib
import io
import importlib
import importlib.util
import ipaddress
import logging
import mimetypes
import os
import socket
import sys
import tempfile
import threading
import time
import uuid
from collections.abc import Mapping, Sequence, Callable
from dataclasses import dataclass, field
from pathlib import Path
from types import ModuleType
from typing import TYPE_CHECKING, Any, Literal, Optional, cast
from urllib.parse import urlparse

import chromadb
import pdfplumber
import requests
from bs4 import BeautifulSoup
from chromadb.config import Settings
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, Query, Request, Response, UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.embeddings import OllamaEmbeddings
from langchain_google_community import GoogleDriveLoader
from prometheus_client import CONTENT_TYPE_LATEST
from pydantic import AliasChoices, BaseModel, ConfigDict, Field

try:
    from .drive_sync import DriveSyncManager
except (ImportError, ValueError):
    from drive_sync import DriveSyncManager  # type: ignore[no-redef]

try:
    from .collection_config import CollectionConfigError, resolve_collection
except (ImportError, ValueError):
    from collection_config import CollectionConfigError, resolve_collection  # type: ignore[no-redef]

if TYPE_CHECKING:
    from langchain.schema import Document
else:
    try:
        from langchain.schema import Document
    except ImportError:
        Document = Any  # fallback for type checking

# --- Metrics module loader ---
def _load_metrics_module() -> ModuleType:
    module_name = "src.ingestor.metrics"
    existing = sys.modules.get(module_name)
    if isinstance(existing, ModuleType):
        return existing
    spec = importlib.util.spec_from_file_location(
        module_name,
        Path(__file__).with_name("metrics.py"),
    )
    if spec and spec.loader:
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module
    raise ImportError("Unable to load metrics module")

ingest_metrics: ModuleType = _load_metrics_module()

# --- Configuration ---
CHROMA_HOST = os.getenv("CHROMA_HOST", "localhost")
CHROMA_PORT = int(os.getenv("CHROMA_PORT", "8000"))
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
EMBED_MODEL = os.getenv("EMBED_MODEL", "nomic-embed-text")
COLLECTION_NAME = os.getenv("COLLECTION_NAME", "rag_education")

# ── Legacy Chroma routing ───────────────────────────────────────────
# These names are the current production Chroma collections. They are no longer
# the business source of truth: `configs/rag_collections.yml` defines the Nexus
# target collections, and `legacy_collection_mapping.yml` is the only accepted
# bridge from historical names to Nexus domains.
COLLECTION_MAP: dict[str, str] = {
    "education": "rag_education",
    "web3": "rag_web3",
    "blockchain": "rag_web3",
    "divers": "rag_divers",
    "maths_premiere": "rag_maths_premiere",
    "default": "rag_education",
}

MATHS_PREMIERE_FALLBACK_FILTERS: dict[str, str] = {
    "matiere": "Mathématiques",
    "niveau": "Première",
    "groupe": "Enseignements de spécialité (EDS)",
}


def resolve_collection_name(
    section: str | None = None,
    collection: str | None = None,
    *,
    allow_quarantine: bool = True,
) -> str:
    """Resolve a ChromaDB collection only through the versioned Nexus mapping."""
    try:
        resolution = resolve_collection(
            section=section,
            collection=collection,
            allow_non_retrievable=allow_quarantine,
        )
    except CollectionConfigError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return resolution.physical_collection


def _resolve_search_target(
    client: Any,
    payload: SearchRequest,
) -> tuple[str, dict[str, Any], bool]:
    """Resolve the collection and effective filters for a search request."""
    requested_collection = resolve_collection_name(
        section=payload.section,
        collection=payload.collection or None,
        allow_quarantine=False,
    )
    effective_filters = dict(payload.filters)
    maths_fallback_applied = False
    explicit_collection_override = bool(payload.collection and payload.collection.strip())

    if explicit_collection_override or (payload.section or "").strip().lower() != "maths_premiere":
        return requested_collection, effective_filters, maths_fallback_applied

    dedicated_collection = client.get_or_create_collection(
        name=COLLECTION_MAP["maths_premiere"],
        metadata={"hnsw:space": "cosine"},
    )
    if dedicated_collection.count() > 0:
        return requested_collection, effective_filters, maths_fallback_applied

    effective_filters.update(MATHS_PREMIERE_FALLBACK_FILTERS)
    maths_fallback_applied = True
    return COLLECTION_MAP["education"], effective_filters, maths_fallback_applied

CHROMA_REQUEST_TIMEOUT = float(os.getenv("CHROMA_REQUEST_TIMEOUT", "30"))
OLLAMA_REQUEST_TIMEOUT = float(os.getenv("OLLAMA_REQUEST_TIMEOUT", "30"))
MAX_REMOTE_BYTES = int(os.getenv("MAX_REMOTE_BYTES", str(10 * 1024 * 1024)))
LOCAL_SOURCE_ROOT = Path(
    os.getenv("LOCAL_SOURCE_ROOT", "/data/uploads")).resolve()
ALLOW_UNRESTRICTED_LOCAL = os.getenv(
    "ALLOW_UNRESTRICTED_LOCAL", "false").lower() == "true"
URL_SCHEMES_ALLOWED = {"http", "https"}

INGEST_CHUNK_SIZE = int(os.getenv("INGEST_CHUNK_SIZE", "1000"))
INGEST_CHUNK_OVERLAP = int(os.getenv("INGEST_CHUNK_OVERLAP", "150"))
METRICS_ENABLED = ingest_metrics.METRICS_ENABLED
MULTIMODAL_ENABLED = os.getenv("MULTIMODAL_ENABLED", "false").lower() == "true"
MM_PARSER_TIMEOUT = float(os.getenv("MM_PARSER_TIMEOUT", "30"))
MM_MAX_CHARS_PER_CHUNK = int(os.getenv("MM_MAX_CHARS_PER_CHUNK", "1200"))
MM_CACHE_DIR = os.getenv("MM_CACHE_DIR", "/data/mm-cache")
GOOGLE_DRIVE_TOKEN_PATH = os.getenv("GOOGLE_DRIVE_TOKEN_PATH", "/tmp/google-drive-token.json")
GDRIVE_MAX_DOCS = int(os.getenv("GDRIVE_MAX_DOCS", "0"))

# Keep metrics isolated per module import to avoid duplicate registration in tests.
METRIC_REGISTRY = ingest_metrics.REGISTRY
REQUEST_COUNT = ingest_metrics.REQUEST_COUNT
REQUEST_LATENCY = ingest_metrics.REQUEST_LATENCY
INGEST_RESULT = ingest_metrics.INGEST_RESULT
ingest_requests_total = ingest_metrics.ingest_requests_total

logger = logging.getLogger(__name__)

UPLOAD_FILES_PARAM = File(...)

MEDIA_SOURCE_TYPES = frozenset({"video", "image", "audio"})

# Multimodal adapter: default to safe stubs, then try to override with real impl
class Chunk:
    def __init__(self, *args: Any, **kwargs: Any) -> None:  # pragma: no cover - stub
        self.text: str = ""
        self.modality: str = "unknown"
        self.metadata: dict[str, Any] = {}

    def as_text(self) -> str:  # pragma: no cover - stub
        return getattr(self, "text", "")


def _parse_multimodal_stub(*args: Any, **kwargs: Any) -> Any:  # pragma: no cover - stub
    raise RuntimeError("Multimodal parser not available on this runtime")

# Assign stub by default; successful import below will override this name
parse_multimodal: Callable[..., Any] = _parse_multimodal_stub

# Try package-relative import first, then fallback to local module path
_mm_mod: Optional[ModuleType] = None
try:
    from . import mm_adapter as _pkg_mm  # prefer package-relative when available
    _mm_mod = _pkg_mm
except Exception:  # pragma: no cover - import fallback
    try:
        _mm_mod = importlib.import_module("mm_adapter")
    except Exception:
        _mm_mod = None

if _mm_mod is not None:
    parse_multimodal = getattr(_mm_mod, "parse_multimodal", parse_multimodal)

# Import admin_api from the same package; support both package and script execution
try:
    from . import admin_api as _admin_api_module
except Exception:  # pragma: no cover - flexible fallback
    try:
        _admin_api_module = importlib.import_module("ingestor.admin_api")
    except Exception:
        _admin_api_module = importlib.import_module("admin_api")

admin_api = _admin_api_module


@dataclass
class PreparedBatch:
    ids: list[str]
    documents: list[str]
    metadatas: list[dict[str, str]]
    modality: str


@dataclass
class DriveTaskProgress:
    """Suivi de progression pour une ingestion Google Drive."""
    task_id: str
    folder_id: str
    status: str = "pending"  # pending, scanning, ingesting, done, error
    total_files: int = 0
    processed_files: int = 0
    added_chunks: int = 0
    skipped_files: int = 0
    error_files: int = 0
    current_file: str = ""
    target_collection: str = ""
    file_results: list[dict[str, Any]] = field(default_factory=list)
    error_message: str = ""
    started_at: float = 0.0
    finished_at: float = 0.0


# In-memory store for drive task progress (thread-safe via GIL for simple dict ops)
_drive_tasks: dict[str, DriveTaskProgress] = {}
_drive_tasks_lock = threading.Lock()

app = FastAPI(title="RAG Ingestor API")
app.include_router(admin_api.router)

_DRIVE_SYNC_DB = os.getenv("DRIVE_SYNC_DB_PATH", "/data/drive_sync_state.db")
sync_manager = DriveSyncManager(db_path=_DRIVE_SYNC_DB)


@app.middleware("http")
async def _metrics_middleware(request, call_next):
    start = time.perf_counter()
    code = 500
    try:
        response = await call_next(request)
        code = getattr(response, "status_code", 500)
    except Exception:
        code = 500
        raise
    finally:
        if ingest_metrics.METRICS_ENABLED:
            elapsed = time.perf_counter() - start
            path = request.url.path
            method = request.method
            REQUEST_LATENCY.labels(path=path, method=method).observe(elapsed)
            REQUEST_COUNT.labels(path=path, method=method, code=str(code)).inc()
    return response


def _record_ingest_metrics(ok: bool) -> None:
    if ingest_metrics.METRICS_ENABLED:
        INGEST_RESULT.labels(status="ok" if ok else "fail").inc()


def _record_ingest_outcome(source: str, modality: str, status: str) -> None:
    if not METRICS_ENABLED:
        return
    safe_source = (source or "unknown").strip().lower() or "unknown"
    safe_modality = (modality or "unknown").strip().lower() or "unknown"
    safe_status = (status or "unknown").strip().lower() or "unknown"
    ingest_requests_total.labels(
        source=safe_source, modality=safe_modality, status=safe_status
    ).inc()

# --- Modèle de requête ---


class IngestRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True, extra="ignore")
    source_type: Literal["url", "gdrive_folder", "pdf", "docx", "markdown", "md", "video", "image", "audio", "auto"] = Field(
        alias="sourceType",
        validation_alias=AliasChoices("source_type", "sourceType"),
    )
    source: str = Field(
        alias="sourceUrl",
        validation_alias=AliasChoices("source", "sourceUrl"),
    )
    metadata_hints: dict[str, str] = Field(
        default_factory=dict,
        alias="metadata",
        validation_alias=AliasChoices("hints", "metadata"),
    )


class DriveIngestRequest(BaseModel):
    folder_id: str
    metadata: dict[str, str] = Field(default_factory=dict)


class UrlBatchRequest(BaseModel):
    """Requête d'ingestion par lot d'URLs."""
    model_config = ConfigDict(populate_by_name=True, extra="ignore")
    urls: list[str] = Field(description="Liste d'URLs à ingérer")
    metadata_hints: dict[str, str] = Field(
        default_factory=dict,
        alias="metadata",
        validation_alias=AliasChoices("hints", "metadata"),
    )


class DeduplicationCheckRequest(BaseModel):
    """Requête de vérification de doublons avant ingestion."""
    sources: list[str] = Field(description="Liste de source_path ou URLs à vérifier")
    section: str = Field(default="education", description="Section pour cibler la bonne collection")
    collection: str = Field(default="", description="Collection explicite à utiliser pour la vérification")


class SearchRequest(BaseModel):
    q: str = Field(description="Query text")
    k: int = Field(default=6, ge=1, le=50, description="Number of results")
    include_documents: bool = Field(default=True, description="Include full text in hits")
    score_threshold: float | None = Field(default=None, ge=0.0, description="Maximum accepted Chroma distance")
    collection: str = Field(
        default="",
        description="Target collection name (overrides section-based routing)",
    )
    section: str = Field(
        default="education",
        description="Section (education, web3) — used to route to the correct collection",
    )
    filters: dict[str, Any] = Field(
        default_factory=dict,
        description="Metadata filters (matiere, niveau, groupe, type_ressource, etc.)",
    )


def _validate_upload_mode(source_type: str, mode: str) -> None:
    normalized_source_type = (source_type or "").strip().lower()
    normalized_mode = (mode or "text").strip().lower() or "text"
    if normalized_source_type not in MEDIA_SOURCE_TYPES:
        return
    if normalized_mode != "multimodal":
        raise HTTPException(
            status_code=400,
            detail="Ce type de fichier nécessite mode=multimodal.",
        )
    if not MULTIMODAL_ENABLED:
        raise HTTPException(status_code=400, detail="Multimodal ingest disabled")

# --- Utilitaires ---


def normalize_metadata(d: dict) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for key, value in d.items():
        if value is None or value == "":
            continue
        if isinstance(value, (list, dict)) and not value:
            continue
        normalized[str(key).strip().lower().replace(" ", "_")] = str(value)
    return normalized


def get_content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def get_content_fingerprint(texts: Sequence[str]) -> str:
    parts = [(text or "").strip() for text in texts]
    normalized = [part for part in parts if part]
    if not normalized:
        return ""
    return hashlib.sha256("\n\n".join(normalized).encode("utf-8")).hexdigest()


def _resolve_local_path(raw_path: str) -> Path:
    candidate = Path(raw_path)
    if not candidate.is_absolute():
        candidate = (LOCAL_SOURCE_ROOT / candidate).resolve()
    else:
        candidate = candidate.resolve()
    if not ALLOW_UNRESTRICTED_LOCAL and not str(candidate).startswith(str(LOCAL_SOURCE_ROOT)):
        raise HTTPException(
            status_code=400, detail="Chemin local en dehors de la zone autorisée")
    if not candidate.exists():
        raise HTTPException(status_code=400, detail="Fichier introuvable")
    if not candidate.is_file():
        raise HTTPException(
            status_code=400, detail="Le chemin indiqué n'est pas un fichier")
    return candidate


def _get_client_ip(request: Any) -> str:
    headers = getattr(request, "headers", {}) or {}
    forwarded = headers.get("X-Forwarded-For") or headers.get("x-forwarded-for")
    if isinstance(forwarded, str) and forwarded.strip():
        primary = forwarded.split(",")[0].strip()
        if primary:
            return primary
    client = getattr(request, "client", None)
    host = getattr(client, "host", None)
    if isinstance(host, str) and host:
        return host
    return "127.0.0.1"


def _ip_allowed(ip_str: str, allowlist: str | None) -> bool:
    if not allowlist:
        return True
    try:
        ip_obj = ipaddress.ip_address(ip_str)
    except ValueError:
        return False
    for cidr in allowlist.split(","):
        network = cidr.strip()
        if not network:
            continue
        try:
            if ip_obj in ipaddress.ip_network(network, strict=False):
                return True
        except ValueError:
            continue
    return False


def _enforce_security(request: Any, _req: Any) -> None:
    headers = getattr(request, "headers", {}) or {}
    token_env = os.getenv("INGESTOR_API_TOKEN") or os.getenv("INGEST_AUTH_TOKEN")
    if token_env:
        # Try X-API-Token first, then Authorization (Bearer or raw)
        header_token = headers.get("X-API-Token") or headers.get("x-api-token")
        if not header_token:
            auth = headers.get("Authorization") or headers.get("authorization")
            if isinstance(auth, str) and auth.strip():
                value = auth.strip()
                if value.lower().startswith("bearer "):
                    header_token = value.split(" ", 1)[1].strip()
                else:
                    header_token = value
        if header_token != token_env:
            raise HTTPException(status_code=401, detail="Unauthorized")

    allowlist = os.getenv("INGESTOR_IP_ALLOWLIST")
    if allowlist and not _ip_allowed(_get_client_ip(request), allowlist):
        raise HTTPException(status_code=403, detail="Forbidden")


def _require_api_token_configured() -> str:
    token_env = os.getenv("INGESTOR_API_TOKEN") or os.getenv("INGEST_AUTH_TOKEN")
    if not token_env:
        raise HTTPException(status_code=503, detail="Ingestor API token not configured")
    return token_env


def get_chroma_client() -> Any:
    timeout_seconds = max(1, int(CHROMA_REQUEST_TIMEOUT))
    settings = Settings(
        chroma_server_host=CHROMA_HOST,
        chroma_server_http_port=CHROMA_PORT,
        anonymized_telemetry=False,
        chroma_logservice_request_timeout_seconds=timeout_seconds,
        chroma_sysdb_request_timeout_seconds=timeout_seconds,
        chroma_query_request_timeout_seconds=timeout_seconds,
    )
    return chromadb.HttpClient(host=CHROMA_HOST, port=CHROMA_PORT, settings=settings)


def _load_docx_basic(file_path: str) -> list[Document]:
    import docx
    try:
        d = docx.Document(file_path)
    except Exception as e:
        raise HTTPException(
            status_code=400, detail=f"Impossible de lire le DOCX: {e}") from e
    texts = []
    for p in d.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())
    # (option simple; on pourra enrichir avec les tableaux si besoin)
    content = "\n".join(texts).strip()
    if not content:
        return []
    return [Document(page_content=content, metadata={"source": os.path.basename(file_path)})]


def load_docx(file_path: str) -> list[Document]:
    try:
        from unstructured.partition.docx import partition_docx
    except ImportError:  # pragma: no cover - fallback when optional deps missing
        return _load_docx_basic(file_path)

    try:
        elements = partition_docx(filename=file_path, include_metadata=True)
    except Exception:
        logger.warning("partition_docx failed, falling back to basic DOCX loader", exc_info=True)
        return _load_docx_basic(file_path)

    documents: list[Document] = []
    for element in elements:
        text = getattr(element, "text", "") or ""
        text = text.strip()
        if not text:
            continue
        metadata_dict: dict[str, Any] = {}
        metadata_obj = getattr(element, "metadata", None)
        if metadata_obj is not None:
            try:
                raw_meta = metadata_obj.to_dict()
            except AttributeError:
                raw_meta = dict(metadata_obj) if isinstance(metadata_obj, dict) else {}
            for key, value in (raw_meta or {}).items():
                if value in (None, "", [], {}):
                    continue
                metadata_dict[str(key)] = str(value)
        metadata_dict.setdefault("source", os.path.basename(file_path))
        metadata_dict.setdefault(
            "mime_type",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        documents.append(Document(page_content=text, metadata=metadata_dict))

    if not documents:
        return _load_docx_basic(file_path)
    return documents


def _load_docx_documents_from_bytes(content_bytes: bytes, source_name: str) -> list[Document]:
    """Persist Drive DOCX bytes briefly to reuse the existing DOCX loader pipeline."""
    if not content_bytes:
        return []

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir="/tmp") as tmp_docx:
        tmp_docx.write(content_bytes)
        tmp_path = tmp_docx.name
    try:
        documents = load_docx(tmp_path)
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    for document in documents:
        if not document.metadata.get("source"):
            document.metadata["source"] = source_name
    return documents


class TimedOllamaEmbeddings(OllamaEmbeddings):
    """Ollama embeddings client with explicit network timeout."""

    request_timeout: float = OLLAMA_REQUEST_TIMEOUT

    def _process_emb_response(self, input: str) -> list[float]:
        headers = {
            "Content-Type": "application/json",
            **(self.headers or {}),
        }

        try:
            res = requests.post(
                f"{self.base_url}/api/embeddings",
                headers=headers,
                json={"model": self.model, "prompt": input, **self._default_params},
                timeout=self.request_timeout,
            )
        except requests.exceptions.RequestException as e:
            raise ValueError(f"Error raised by inference endpoint: {e}") from e

        if res.status_code != 200:
            raise ValueError(
                f"Error raised by inference API HTTP code: {res.status_code}, {res.text}"
            )
        try:
            t = res.json()
            return cast(list[float], t["embedding"])
        except requests.exceptions.JSONDecodeError as e:
            raise ValueError(
                f"Error raised by inference API: {e}.\nResponse: {res.text}"
            ) from e


def load_markdown(file_path: Path) -> list[Document]:
    try:
        raw_text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raw_text = file_path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Impossible de lire le fichier Markdown: {exc}") from exc

    text = raw_text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Fichier Markdown vide")

    metadata = {"source": str(file_path), "mime_type": "text/markdown"}
    return [Document(page_content=text, metadata=metadata)]


def _ocr_pdf_bytes(pdf_bytes: bytes, max_pages: int = 5) -> list[str]:
    """OCR a PDF by rendering a bounded number of pages to images first."""
    if not pdf_bytes:
        return []

    try:
        import subprocess

        import pytesseract
        from PIL import Image
    except ImportError:
        return []

    texts: list[str] = []
    with tempfile.TemporaryDirectory(dir="/tmp") as tmpdir:
        pdf_path = Path(tmpdir) / "input.pdf"
        pdf_path.write_bytes(pdf_bytes)
        output_prefix = Path(tmpdir) / "page"
        try:
            subprocess.run(
                [
                    "pdftoppm",
                    "-png",
                    "-r",
                    "200",
                    "-f",
                    "1",
                    "-l",
                    str(max(1, max_pages)),
                    str(pdf_path),
                    str(output_prefix),
                ],
                check=True,
                capture_output=True,
                timeout=120,
            )
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
            logger.warning(f"pdftoppm failed or not installed: {e}")
            return []

        for image_path in sorted(Path(tmpdir).glob("page-*.png")):
            with Image.open(image_path) as image:
                try:
                    text = pytesseract.image_to_string(image, lang="fra+eng")
                except Exception:
                    text = pytesseract.image_to_string(image, lang="eng")
            cleaned = (text or "").strip()
            if cleaned:
                texts.append(cleaned)
    return texts


def _extract_pdf_documents_from_bytes(pdf_bytes: bytes, source_name: str) -> list[Document]:
    """Extract per-page PDF documents, using OCR only when text extraction is empty."""
    docs: list[Document] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf_doc:
        for page_index, page in enumerate(pdf_doc.pages):
            text = (page.extract_text() or "").strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"source": source_name, "page": page_index},
                    )
                )
    if docs:
        return docs

    for page_index, text in enumerate(_ocr_pdf_bytes(pdf_bytes)):
        docs.append(
            Document(
                page_content=text,
                metadata={"source": source_name, "page": page_index, "ocr": "true"},
            )
        )
    return docs


def _download_drive_file_bytes(file_id: str) -> bytes:
    from google.oauth2 import service_account as _sa
    from googleapiclient.discovery import build as _build

    creds_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "")
    creds = _sa.Credentials.from_service_account_file(
        creds_path,
        scopes=["https://www.googleapis.com/auth/drive.readonly"],
    )
    svc = _build("drive", "v3", credentials=creds, cache_discovery=False)
    return cast(bytes, svc.files().get_media(fileId=file_id).execute())


def _documents_have_text(docs: list[Any]) -> bool:
    """Return True when at least one loaded document has non-empty text content."""
    return any((getattr(doc, "page_content", "") or "").strip() for doc in docs)


def _is_pdf_drive_file(file_name: str, mime_type: str) -> bool:
    return str(file_name or "").lower().endswith(".pdf") or mime_type == "application/pdf"


def _is_docx_drive_file(file_name: str, mime_type: str) -> bool:
    return str(file_name or "").lower().endswith(".docx") or (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def _is_office_spreadsheet_file(file_name: str, mime_type: str) -> bool:
    file_name_lower = str(file_name or "").lower()
    spreadsheet_mimes = {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }
    return file_name_lower.endswith((".xlsx", ".xls")) or mime_type in spreadsheet_mimes


def _validate_remote_url(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme.lower() not in URL_SCHEMES_ALLOWED:
        raise HTTPException(
            status_code=400, detail="Schéma d'URL non autorisé")
    if not parsed.hostname:
        raise HTTPException(
            status_code=400, detail="URL invalide")
    try:
        addr_info = socket.getaddrinfo(parsed.hostname, parsed.port or (
            443 if parsed.scheme == "https" else 80))
    except socket.gaierror as exc:
        raise HTTPException(
            status_code=400, detail=f"Résolution DNS impossible: {exc}") from exc
    for entry in addr_info:
        ip = ipaddress.ip_address(entry[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
            raise HTTPException(
                status_code=400, detail="URL interne non autorisée")


def _download_to_temp(url: str, suffix: str) -> Path:
    _validate_remote_url(url)
    headers = {"User-Agent": os.getenv("USER_AGENT", "rag-local-ingestor/1.0")}
    try:
        with requests.get(url, timeout=30, stream=True, headers=headers) as response:
            response.raise_for_status()
            total = 0
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_file:
                for chunk in response.iter_content(chunk_size=8192):
                    if not chunk:
                        continue
                    total += len(chunk)
                    if total > MAX_REMOTE_BYTES:
                        raise HTTPException(
                            status_code=400, detail="Fichier distant trop volumineux")
                    tmp_file.write(chunk)
                return Path(tmp_file.name)
    except HTTPException:
        raise
    except requests.RequestException as exc:
        raise HTTPException(
            status_code=400, detail=f"Téléchargement impossible: {exc}") from exc


def _fetch_remote_text(url: str) -> tuple[str, str]:
    _validate_remote_url(url)
    headers = {"User-Agent": os.getenv("USER_AGENT", "rag-local-ingestor/1.0")}
    try:
        with requests.get(url, timeout=30, allow_redirects=True, stream=True, headers=headers) as response:
            if response.history:
                for hop in response.history:
                    _validate_remote_url(hop.url)
            _validate_remote_url(response.url)
            declared_length = response.headers.get("content-length")
            if declared_length and int(declared_length) > MAX_REMOTE_BYTES:
                raise HTTPException(
                    status_code=400, detail="Réponse distante trop volumineuse")
            chunks: list[bytes] = []
            total = 0
            for chunk in response.iter_content(chunk_size=8192):
                if not chunk:
                    continue
                total += len(chunk)
                if total > MAX_REMOTE_BYTES:
                    raise HTTPException(
                        status_code=400, detail="Réponse distante trop volumineuse")
                chunks.append(chunk)
            encoding = response.encoding or "utf-8"
            text = b"".join(chunks).decode(encoding, errors="ignore")
            if not text.strip():
                raise HTTPException(
                    status_code=400, detail="Aucun contenu exploitable sur la page")
            return response.url, text
    except HTTPException:
        raise
    except requests.RequestException as exc:
        raise HTTPException(
            status_code=400, detail=f"Téléchargement impossible: {exc}") from exc


def load_from_url(url: str) -> list[Document]:
    if url.lower().endswith(".pdf"):
        tmp_path = _download_to_temp(url, suffix=".pdf")
        try:
            return PyPDFLoader(str(tmp_path)).load()
        finally:
            try:
                tmp_path.unlink()
            except OSError:
                pass
    final_url, text = _fetch_remote_text(url)
    soup = BeautifulSoup(text, "html.parser")
    text = soup.get_text("\n", strip=True)
    if not text:
        raise HTTPException(
            status_code=400, detail="Aucun contenu exploitable sur la page")
    return [Document(page_content=text, metadata={"source": final_url})]


def _load_source_documents(req: IngestRequest) -> list[Document]:
    if req.source_type == "url":
        return load_from_url(req.source)
    if req.source_type == "gdrive_folder":
        loader_kwargs: dict[str, Any] = {"folder_id": req.source, "recursive": True}
        loader_kwargs["supports_all_drives"] = True
        loader_kwargs["export_mime_types"] = {
            "application/vnd.google-apps.document": "text/plain",
            "application/vnd.google-apps.spreadsheet": "text/csv",
            "application/vnd.google-apps.presentation": "text/plain",
        }
        service_account_raw = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        if service_account_raw:
            service_account_path = Path(service_account_raw)
            if not service_account_path.exists():
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Configuration Google Drive invalide: le fichier de clé de service "
                        "spécifié par GOOGLE_APPLICATION_CREDENTIALS est introuvable."
                    ),
                )
            loader_kwargs["service_account_key"] = service_account_path
            loader_kwargs["credentials_path"] = service_account_path
        else:
            default_credentials = Path.home() / ".credentials" / "credentials.json"
            if default_credentials.exists():
                loader_kwargs["credentials_path"] = default_credentials
            else:
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Identification Google Drive manquante: définissez GOOGLE_APPLICATION_CREDENTIALS "
                        "ou placez un credentials.json valide dans ~/.credentials/."
                    ),
                )

        token_path = Path(GOOGLE_DRIVE_TOKEN_PATH)
        try:
            token_path.parent.mkdir(parents=True, exist_ok=True)
        except OSError as exc:  # pragma: no cover - defensive guard
            raise HTTPException(
                status_code=500,
                detail=f"Impossible de préparer le répertoire du token Google Drive: {exc}",
            ) from exc
        loader_kwargs["token_path"] = token_path

        # Fast path: when a limiter is configured, pre-list a few file ids and load only those.
        limit = int(GDRIVE_MAX_DOCS)
        file_ids: list[str] = []
        if limit > 0:
            try:
                # Import inside the branch to avoid hard dependency at import time.
                from google.oauth2 import service_account as _sa
                from googleapiclient.discovery import build as _build
                creds = _sa.Credentials.from_service_account_file(str(loader_kwargs.get("credentials_path", service_account_raw)))
                svc = _build("drive", "v3", credentials=creds, cache_discovery=False)

                def _list_children(parent_id: str, q_extra: str, page_size: int = 10) -> list[dict[str, Any]]:
                    q = f"'{parent_id}' in parents and trashed=false {q_extra}"
                    resp = svc.files().list(
                        q=q,
                        pageSize=page_size,
                        fields="files(id,name,mimeType)",
                        includeItemsFromAllDrives=True,
                        supportsAllDrives=True,
                    ).execute()
                    return list(resp.get("files", []))

                # Shallow BFS up to depth 2 to find up to `limit` non-folder files quickly.
                queue: list[tuple[str, int]] = [(req.source, 0)]
                seen: set[str] = {req.source}
                while queue and len(file_ids) < limit:
                    current, depth = queue.pop(0)
                    try:
                        files = _list_children(current, "and mimeType != 'application/vnd.google-apps.folder'", page_size=max(5, limit))
                        for f in files:
                            if len(file_ids) >= limit:
                                break
                            fid = str(f.get("id", "") or "")
                            if fid:
                                file_ids.append(fid)
                        if len(file_ids) >= limit or depth >= 2:
                            continue
                        subs = _list_children(current, "and mimeType = 'application/vnd.google-apps.folder'", page_size=5)
                        for sf in subs:
                            sid = str(sf.get("id", "") or "")
                            if sid and sid not in seen:
                                seen.add(sid)
                                queue.append((sid, depth + 1))
                    except Exception:
                        # Ignore listing errors at this stage; we'll fall back to loader recursion.
                        continue

                if file_ids:
                    # Switch to file_ids mode for faster, bounded loading
                    loader_kwargs.pop("folder_id", None)
                    loader_kwargs.pop("recursive", None)
                    loader_kwargs["file_ids"] = file_ids
            except Exception:
                # non-fatal; proceed with regular folder traversal
                pass

        try:
            loader = GoogleDriveLoader(**loader_kwargs)
        except ValueError as exc:
            raise HTTPException(
                status_code=500,
                detail=f"Configuration Google Drive invalide: {exc}",
            ) from exc

        docs: list[Document] = []
        try:
            _lazy = getattr(loader, "lazy_load", None)
            if callable(_lazy):
                _result = _lazy()
                _iterable: list[Any] = list(_result) if hasattr(_result, "__iter__") else []
                for _d in _iterable:
                    try:
                        if _d and getattr(_d, "page_content", "").strip():
                            docs.append(_d)
                            if limit > 0 and len(docs) >= limit:
                                break
                    except Exception:
                        continue
            else:
                docs = loader.load()
                if limit > 0 and len(docs) > limit:
                    docs = docs[:limit]
        except Exception as _e:
            try:
                docs = []
                for _d in loader.lazy_load():
                    try:
                        if _d and getattr(_d, "page_content", "").strip():
                            docs.append(_d)
                            if limit > 0 and len(docs) >= limit:
                                break
                    except Exception:
                        continue
            except Exception as _e2:
                raise HTTPException(status_code=500, detail=f"Echec chargement Google Drive: {_e2}") from _e2
        # Ne pas échouer si aucun document lisible: laissez l'API retourner added:0
        return docs
    if req.source_type == "pdf":
        path = _resolve_local_path(req.source)
        return PyPDFLoader(str(path)).load()
    if req.source_type == "docx":
        path = _resolve_local_path(req.source)
        return load_docx(str(path))
    if req.source_type in {"markdown", "md"}:
        path = _resolve_local_path(req.source)
        return load_markdown(path)
    if req.source_type in {"video", "image", "audio"}:
        raise HTTPException(
            status_code=400,
            detail=(
                f"Ingestion {req.source_type} disponible uniquement en mode multimodal "
                "(mode=multimodal). Utilisez /ingest/upload-files pour un traitement automatique."
            ),
        )
    if req.source_type == "auto":
        path = _resolve_local_path(req.source)
        ext = path.suffix.lower()
        if ext == ".pdf":
            return PyPDFLoader(str(path)).load()
        if ext in {".docx", ".doc"}:
            return load_docx(str(path))
        if ext in {".md", ".markdown", ".txt", ".csv", ".html", ".htm"}:
            return load_markdown(path)
        # Default: try as text
        return load_markdown(path)
    raise HTTPException(status_code=400, detail=f"source_type non géré: {req.source_type}")


def _prepare_chunks_for_chroma(
    req: IngestRequest,
    docs: list[Document],
    splitter: Optional[RecursiveCharacterTextSplitter] = None,
    extra_metadata: dict[str, Any] | None = None,
) -> PreparedBatch:
    splitter = splitter or RecursiveCharacterTextSplitter(
        chunk_size=INGEST_CHUNK_SIZE, chunk_overlap=INGEST_CHUNK_OVERLAP
    )
    chunks = splitter.split_documents(docs)
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, str]] = []
    modality = "text"
    seen_ids: set[str] = set()
    chunk_number = 0

    for chunk in chunks:
        text = (chunk.page_content or "").strip()
        if not text:
            continue
        content_hash = get_content_hash(text)
        if content_hash in seen_ids:
            continue
        seen_ids.add(content_hash)
        chunk_modality = (chunk.metadata or {}).get("modality", "text")
        metadata: dict[str, Any] = {
            "sha256": content_hash,
            "source_type": req.source_type,
            "source": req.source,
            "modality": chunk_modality,
            "chunk_index": str(chunk_number),
        }
        metadata.update(chunk.metadata or {})
        if extra_metadata:
            metadata.update(extra_metadata)
        metadata.update(req.metadata_hints or {})
        normalized = normalize_metadata(metadata)

        ids.append(content_hash)
        documents.append(text)
        metadatas.append(normalized)
        modality = normalized.get("modality", modality)
        chunk_number += 1

    if not ids:
        modality = "text"
    return PreparedBatch(ids=ids, documents=documents, metadatas=metadatas, modality=modality)


def _prepare_multimodal_chunks(req: IngestRequest, chunks: list[Chunk]) -> PreparedBatch:
    ids: list[str] = []
    documents: list[str] = []
    metadatas: list[dict[str, str]] = []
    modality_counts: dict[str, int] = {}
    seen_ids: set[str] = set()

    for chunk in chunks:
        text = chunk.as_text() if hasattr(chunk, "as_text") else (chunk.text or "")
        text = (text or "").strip()
        if not text:
            continue
        content_hash = get_content_hash(text)
        if content_hash in seen_ids:
            continue
        seen_ids.add(content_hash)
        chunk_modality = (chunk.modality or "unknown").strip().lower() or "unknown"
        metadata: dict[str, Any] = {
            "sha256": content_hash,
            "source": req.source,
            "source_type": req.source_type,
            "modality": chunk_modality,
        }
        metadata.update(getattr(chunk, "metadata", {}) or {})
        metadata.update(req.metadata_hints or {})
        normalized = normalize_metadata(metadata)

        ids.append(content_hash)
        documents.append(text)
        metadatas.append(normalized)

        key = normalized.get("modality", chunk_modality)
        modality_counts[key] = modality_counts.get(key, 0) + 1

    if modality_counts:
        dominant = max(modality_counts.items(), key=lambda item: item[1])[0]
    else:
        dominant = "unknown"
    return PreparedBatch(ids=ids, documents=documents, metadatas=metadatas, modality=dominant)


def _prepare_multimodal_ingest(req: IngestRequest) -> PreparedBatch:
    if not MULTIMODAL_ENABLED:
        raise HTTPException(status_code=400, detail="Multimodal ingest disabled")
    path = _resolve_local_path(req.source)
    mime, _ = mimetypes.guess_type(path.name)
    with path.open("rb") as handle:
        chunk_iter = parse_multimodal(
            handle,
            filename=path.name,
            mime=mime or "application/octet-stream",
            timeout_s=MM_PARSER_TIMEOUT,
            max_chars_per_chunk=MM_MAX_CHARS_PER_CHUNK,
            cache_dir=MM_CACHE_DIR,
        )
        chunk_list = list(chunk_iter)
    return _prepare_multimodal_chunks(req, chunk_list)

# --- Endpoint ---


def _index_batch(
    prepared: PreparedBatch,
    req_source_type: str,
    modality_label: str,
    collection_name: str | None = None,
) -> dict[str, Any]:
    target_collection = collection_name or COLLECTION_NAME
    if not prepared.ids:
        _record_ingest_metrics(True)
        _record_ingest_outcome(req_source_type, modality_label, "empty")
        return {"status": "ok", "message": "Aucun contenu éligible à l'ingestion."}

    dedup_ids: list[str] = []
    dedup_documents: list[str] = []
    dedup_metadatas: list[dict[str, str]] = []
    seen_ids: set[str] = set()
    batch_skipped = 0
    for chunk_id, document, metadata in zip(prepared.ids, prepared.documents, prepared.metadatas, strict=False):
        if chunk_id in seen_ids:
            batch_skipped += 1
            continue
        seen_ids.add(chunk_id)
        dedup_ids.append(chunk_id)
        dedup_documents.append(document)
        dedup_metadatas.append(metadata)

    prepared = PreparedBatch(
        ids=dedup_ids,
        documents=dedup_documents,
        metadatas=dedup_metadatas,
        modality=prepared.modality,
    )

    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(
            name=target_collection, metadata={"hnsw:space": "cosine"}
        )

        existing = collection.get(ids=prepared.ids) or {}
        existing_ids = set(existing.get("ids", []))

        to_add_idx = [i for i, chunk_id in enumerate(prepared.ids) if chunk_id not in existing_ids]
        if not to_add_idx:
            _record_ingest_metrics(True)
            _record_ingest_outcome(req_source_type, modality_label, "skipped")
            return {"status": "ok", "added": 0, "skipped": len(existing_ids) + batch_skipped, "collection": target_collection}

        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        docs_to_add = [prepared.documents[i] for i in to_add_idx]
        ids_to_add = [prepared.ids[i] for i in to_add_idx]
        meta_to_add = [prepared.metadatas[i] for i in to_add_idx]
        try:
            embs_to_add = emb.embed_documents(docs_to_add)
        except ValueError as exc:
            message = str(exc)
            if "HTTP code: 404" in message:
                logger.warning(
                    "Ollama embeddings endpoint returned 404 for model '%s'", EMBED_MODEL
                )
                raise HTTPException(
                    status_code=503,
                    detail=(
                        f"Embedding model '{EMBED_MODEL}' is not available on the Ollama backend. "
                        "Pull the model or adjust EMBED_MODEL before retrying."
                    ),
                ) from exc
            logger.exception("Embedding provider raised ValueError")
            raise
        except Exception:  # pragma: no cover - defensive logging
            logger.exception("Unexpected failure while requesting embeddings")
            raise

        meta_mappings = cast(list[Mapping[str, Any]], meta_to_add)
        embeddings_seq = cast(list[Sequence[float]], embs_to_add)
        collection.add(
            documents=docs_to_add,
            ids=ids_to_add,
            metadatas=meta_mappings,
            embeddings=embeddings_seq,
        )
        _record_ingest_metrics(True)
        _record_ingest_outcome(req_source_type, modality_label, "success")
        return {
            "status": "ok",
            "added": len(ids_to_add),
            "skipped": (len(prepared.ids) - len(ids_to_add)) + batch_skipped,
            "collection": target_collection,
        }
    except HTTPException as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req_source_type, modality_label, f"http_{exc.status_code}")
        raise
    except Exception as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req_source_type, modality_label, "error")
        raise HTTPException(
            status_code=500, detail=f"Erreur d'ingestion dans ChromaDB: {exc}"
        ) from exc


def background_drive_ingest(folder_id: str, metadata: dict, task_id: str) -> None:
    """Ingestion Google Drive en arrière-plan avec suivi de progression."""
    task = _drive_tasks.get(task_id)
    if not task:
        logger.error(f"Drive task {task_id} not found in store")
        return

    task.status = "scanning"
    task.started_at = time.time()
    target_col = resolve_collection_name(section=metadata.get("section"), collection=metadata.get("collection"))
    task.target_collection = target_col

    logger.info(f"[{task_id}] Starting drive ingest for folder {folder_id} → {target_col}")

    try:
        # Vérification d'accès explicite avant de démarrer — lève PermissionError si le SA n'a pas accès
        try:
            sync_manager.verify_folder_access(folder_id)
        except PermissionError as e:
            logger.error(f"[{task_id}] Drive access denied: {e}")
            task.status = "error"
            task.error_message = str(e)
            task.finished_at = time.time()
            return
        except Exception as e:
            logger.error(f"[{task_id}] Drive access check failed: {e}")
            task.status = "error"
            task.error_message = f"Erreur vérification accès Drive: {e}"
            task.finished_at = time.time()
            return

        updates = sync_manager.list_updates(folder_id, collection_name=target_col)
        if not updates:
            logger.info(f"[{task_id}] Aucune mise à jour pour le dossier Drive.")
            task.status = "done"
            task.finished_at = time.time()
            return

        task.total_files = len(updates)
        task.status = "ingesting"

        # Setup credentials once
        loader_kwargs_base: dict[str, Any] = {}
        loader_kwargs_base["supports_all_drives"] = True
        loader_kwargs_base["export_mime_types"] = {
            "application/vnd.google-apps.document": "text/plain",
            "application/vnd.google-apps.spreadsheet": "text/csv",
            "application/vnd.google-apps.presentation": "text/plain",
        }
        service_account_raw = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
        if service_account_raw:
            service_account_path = Path(service_account_raw)
            if service_account_path.exists():
                loader_kwargs_base["service_account_key"] = service_account_path
                loader_kwargs_base["credentials_path"] = service_account_path
        else:
            default_credentials = Path.home() / ".credentials" / "credentials.json"
            if default_credentials.exists():
                loader_kwargs_base["credentials_path"] = default_credentials

        token_path = Path(GOOGLE_DRIVE_TOKEN_PATH)
        try:
            token_path.parent.mkdir(parents=True, exist_ok=True)
        except OSError:
            pass
        loader_kwargs_base["token_path"] = token_path

        for file_meta in updates:
            fid = file_meta.get("id", "?")
            fname = file_meta.get("name", "?")
            mime_type = str(file_meta.get("mimeType") or "")
            task.current_file = fname

            fname_lower = fname.lower()
            unsupported_exts = (
                ".webm", ".mkv", ".mp4", ".avi", ".mov", ".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac",
                ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg", ".tif", ".tiff", ".ico",
                ".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz",
                ".ppt", ".pptx", ".pps", ".ppsx", ".php",
                ".html", ".htm", ".txt",
                # LaTeX auxiliary/compiled files
                ".out", ".snm", ".aux", ".log", ".toc", ".bbl", ".blg", ".nav", ".vrb", ".fls", ".fdb_latexmk",
                # PostScript (incompatible avec pypdf)
                ".ps", ".eps",
            )
            unsupported_mimes = (
                "video/", "audio/", "image/",
                "application/zip", "application/x-zip", "application/x-zip-compressed",
                "application/x-gzip", "application/gzip", "application/x-tar",
                "application/x-rar", "application/vnd.rar", "application/x-7z-compressed",
                "application/vnd.ms-powerpoint",
                "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "application/vnd.openxmlformats-officedocument.presentationml.slideshow",
                "text/html",
                "application/postscript",
            )
            if mime_type.startswith(unsupported_mimes) or fname_lower.endswith(unsupported_exts):
                logger.info(f"[{task_id}] Skipping unsupported Drive file {fid} ({fname}) [{mime_type}]")
                task.skipped_files += 1
                task.processed_files += 1
                task.file_results.append({
                    "name": fname,
                    "status": "unsupported",
                    "added": 0,
                    "skipped": 1,
                    "detail": f"Type Google Drive non supporté pour ce flux: {mime_type or fname}",
                })
                continue

            if _is_office_spreadsheet_file(fname, mime_type):
                logger.info(f"[{task_id}] Skipping unsupported spreadsheet Drive file {fid} ({fname}) [{mime_type}]")
                task.skipped_files += 1
                task.processed_files += 1
                task.file_results.append({
                    "name": fname,
                    "status": "unsupported",
                    "added": 0,
                    "skipped": 1,
                    "detail": "Type spreadsheet Office non supporté pour ce flux.",
                })
                continue

            # .doc ancien → conversion libreoffice avant chargement
            is_legacy_doc = fname_lower.endswith(".doc") or mime_type == "application/msword"
            is_modern_docx = _is_docx_drive_file(fname, mime_type)

            try:
                logger.info(f"[{task_id}] Processing file {fid} ({fname})")

                # Cas 1 : Google Docs/Sheets/Slides → export texte natif via Drive API
                google_workspace_export: dict[str, str] = {
                    "application/vnd.google-apps.document": "text/plain",
                    "application/vnd.google-apps.spreadsheet": "text/csv",
                    "application/vnd.google-apps.presentation": "text/plain",
                }
                docs: list[Any] = []
                if mime_type in google_workspace_export:
                    try:
                        from google.oauth2 import service_account as _sa
                        from googleapiclient.discovery import build as _build
                        from langchain.schema import Document as _Doc
                        _creds_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "")
                        _creds = _sa.Credentials.from_service_account_file(
                            _creds_path, scopes=["https://www.googleapis.com/auth/drive.readonly"]
                        )
                        _svc = _build("drive", "v3", credentials=_creds, cache_discovery=False)
                        export_mime = google_workspace_export[mime_type]
                        content_bytes = _svc.files().export(
                            fileId=fid, mimeType=export_mime
                        ).execute()
                        text = content_bytes.decode("utf-8", errors="replace") if isinstance(content_bytes, bytes) else str(content_bytes)
                        if text.strip():
                            docs = [_Doc(page_content=text, metadata={"source": fname})]
                        logger.info(f"[{task_id}] Google Workspace export OK: {fname} ({export_mime})")
                    except Exception as _e:
                        logger.warning(f"[{task_id}] Google Workspace export failed for {fid}: {_e}")
                        # Fallback: tenter GoogleDriveLoader
                        docs = []

                # Cas 2 : .docx moderne → téléchargement direct → pipeline DOCX
                elif is_modern_docx:
                    try:
                        content_bytes = _download_drive_file_bytes(fid)
                        docs = _load_docx_documents_from_bytes(content_bytes, fname)
                        logger.info(f"[{task_id}] .docx loaded directly: {fname}")
                    except Exception as _e:
                        logger.warning(f"[{task_id}] .docx load failed for {fid}: {_e}")
                        docs = []

                # Cas 3 : .doc ancien → libreoffice → .docx → python-docx
                elif is_legacy_doc:
                    try:
                        import subprocess
                        content_bytes = _download_drive_file_bytes(fid)
                        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False, dir="/tmp") as tf:
                            tf.write(content_bytes)
                            tmp_doc = tf.name
                        conv_res = subprocess.run(
                            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", "/tmp", tmp_doc],
                            capture_output=True, timeout=60
                        )
                        tmp_docx = tmp_doc.replace(".doc", ".docx")
                        if conv_res.returncode == 0 and Path(tmp_docx).exists():
                            docs = load_docx(tmp_docx)
                            Path(tmp_doc).unlink(missing_ok=True)
                            Path(tmp_docx).unlink(missing_ok=True)
                            logger.info(f"[{task_id}] .doc converted via libreoffice: {fname}")
                        else:
                            Path(tmp_doc).unlink(missing_ok=True)
                            raise RuntimeError(f"libreoffice conversion failed: {conv_res.stderr.decode()[:200]}")
                    except Exception as _e:
                        logger.warning(f"[{task_id}] .doc conversion failed for {fid}: {_e}")
                        docs = []

                # Cas 4 : chargement standard via GoogleDriveLoader
                if not docs and mime_type not in google_workspace_export and not is_legacy_doc and not is_modern_docx:
                    kwargs = loader_kwargs_base.copy()
                    kwargs["file_ids"] = [fid]
                    try:
                        loader = GoogleDriveLoader(**kwargs)
                        docs = loader.load()
                        if not _documents_have_text(docs) and _is_pdf_drive_file(fname, mime_type):
                            pdf_bytes = _download_drive_file_bytes(fid)
                            docs = _extract_pdf_documents_from_bytes(pdf_bytes, fname)
                            if docs:
                                logger.info(f"[{task_id}] pdf/OCR fallback OK after empty loader result: {fname}")
                    except Exception as e:
                        error_detail = str(e)
                        # Fallback pdfplumber pour PDFs corrompus (EOF marker not found)
                        pdf_error_markers = ("eof marker", "startxref", "cannot read an empty", "trailer", "malformed pdf", "invalid pdf", "file is not a pdf")
                        if _is_pdf_drive_file(fname, mime_type) and any(m in error_detail.lower() for m in pdf_error_markers):
                            try:
                                pdf_bytes = _download_drive_file_bytes(fid)
                                docs = _extract_pdf_documents_from_bytes(pdf_bytes, fname)
                                if docs:
                                    logger.info(f"[{task_id}] pdf/OCR fallback OK: {fname} ({len(docs)} pages)")
                                else:
                                    raise RuntimeError("pdfplumber/OCR extracted no text")
                            except Exception as _fe:
                                logger.warning(f"[{task_id}] pdfplumber fallback failed for {fid}: {_fe}")
                                error_detail2 = f"{error_detail} | pdfplumber: {_fe}"
                                task.processed_files += 1
                                task.skipped_files += 1
                                task.file_results.append({"name": fname, "status": "invalid", "added": 0, "skipped": 1, "detail": error_detail2})
                                continue
                        else:
                            logger.warning(f"[{task_id}] Failed to load file {fid}: {error_detail}")
                            task.processed_files += 1
                            task.error_files += 1
                            task.file_results.append({"name": fname, "status": "error", "detail": error_detail})
                            continue

                if not docs:
                    logger.warning(f"[{task_id}] No content loaded for file {fid}")
                    task.skipped_files += 1
                    task.processed_files += 1
                    task.file_results.append({"name": fname, "status": "empty", "added": 0})
                    continue

                content_fingerprint = get_content_fingerprint([(doc.page_content or "") for doc in docs])
                if sync_manager.is_unchanged(
                    file_meta,
                    content_fingerprint,
                    collection_name=target_col,
                ):
                    logger.info(f"[{task_id}] Skipping unchanged file {fid} ({fname})")
                    task.skipped_files += 1
                    task.processed_files += 1
                    task.file_results.append({
                        "name": fname,
                        "status": "duplicate",
                        "added": 0,
                        "skipped": 1,
                        "detail": "Contenu inchangé (empreinte identique).",
                    })
                    continue

                # Prepare chunks — inject Drive-specific metadata per file
                req = IngestRequest(
                    sourceType="gdrive_folder",
                    sourceUrl=folder_id,
                    metadata=metadata,
                )
                drive_extra: dict[str, Any] = {
                    "drive_file_id": fid,
                    "drive_file_name": fname,
                    "drive_folder_id": folder_id,
                    "mime_type": mime_type,
                }
                prepared = _prepare_chunks_for_chroma(req, docs, extra_metadata=drive_extra)

                if not prepared.ids:
                    task.skipped_files += 1
                    task.processed_files += 1
                    task.file_results.append({"name": fname, "status": "empty", "added": 0})
                    continue

                # Index
                try:
                    batch_res = _index_batch(prepared, "gdrive_folder", "text", collection_name=target_col)
                    added = batch_res.get("added", 0)
                    skipped = batch_res.get("skipped", 0)
                    task.added_chunks += added
                    if added == 0 and skipped > 0:
                        task.skipped_files += 1
                        task.file_results.append({"name": fname, "status": "duplicate", "added": 0, "skipped": skipped})
                    else:
                        task.file_results.append({"name": fname, "status": "ok", "added": added, "skipped": skipped})
                except Exception as e:
                    logger.error(f"[{task_id}] Indexing failed for {fid}: {e}")
                    task.error_files += 1
                    task.file_results.append({"name": fname, "status": "error", "detail": str(e)})
                    task.processed_files += 1
                    continue

                # Mark as ingested
                sync_manager.mark_as_ingested(
                    file_meta,
                    content_fingerprint=content_fingerprint,
                    collection_name=target_col,
                )

            except Exception as e:
                logger.error(f"[{task_id}] Error processing file {fid}: {e}")
                task.error_files += 1
                task.file_results.append({"name": fname, "status": "error", "detail": str(e)})

            task.processed_files += 1

        task.status = "done"
        task.current_file = ""
        task.finished_at = time.time()
        logger.info(
            f"[{task_id}] Drive ingest done: {task.processed_files}/{task.total_files} files, "
            f"{task.added_chunks} chunks added → {target_col}"
        )

    except Exception as e:
        logger.error(f"[{task_id}] Background ingest failed: {e}")
        task.status = "error"
        task.error_message = str(e)
        task.finished_at = time.time()


@app.post("/ingest/drive")
def ingest_drive(
    req: DriveIngestRequest,
    background_tasks: BackgroundTasks,
    request: Request,
):
    """Lance une ingestion Google Drive et retourne un task_id pour suivi de progression."""
    _require_api_token_configured()
    _enforce_security(request, req)

    task_id = uuid.uuid4().hex[:12]
    target_col = resolve_collection_name(section=req.metadata.get("section"), collection=req.metadata.get("collection"))
    task = DriveTaskProgress(task_id=task_id, folder_id=req.folder_id, target_collection=target_col)

    with _drive_tasks_lock:
        _drive_tasks[task_id] = task

    background_tasks.add_task(background_drive_ingest, req.folder_id, req.metadata, task_id)

    return {
        "status": "accepted",
        "task_id": task_id,
        "target_collection": target_col,
        "message": f"Ingestion Drive démarrée (tâche {task_id})",
    }


@app.get("/ingest/drive/status/{task_id}")
def drive_ingest_status(task_id: str, request: Request) -> dict[str, Any]:
    """Retourne la progression d'une ingestion Google Drive."""
    _require_api_token_configured()
    _enforce_security(request, None)

    task = _drive_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail=f"Tâche {task_id} introuvable")

    elapsed = 0.0
    if task.started_at > 0:
        end = task.finished_at if task.finished_at > 0 else time.time()
        elapsed = round(end - task.started_at, 1)

    progress_pct = 0
    if task.total_files > 0:
        progress_pct = round((task.processed_files / task.total_files) * 100)

    return {
        "task_id": task.task_id,
        "folder_id": task.folder_id,
        "status": task.status,
        "target_collection": task.target_collection,
        "total_files": task.total_files,
        "processed_files": task.processed_files,
        "added_chunks": task.added_chunks,
        "skipped_files": task.skipped_files,
        "error_files": task.error_files,
        "current_file": task.current_file,
        "progress_pct": progress_pct,
        "elapsed_seconds": elapsed,
        "error_message": task.error_message,
        "file_results": task.file_results,
    }


@app.post("/ingest")
def ingest_data(
    req: IngestRequest,
    request: Request,
    mode: str = Query(default="text"),
):
    modality_label = "unknown"
    mode_normalized = (mode or "text").strip().lower() or "text"

    try:
        _require_api_token_configured()
        _enforce_security(request, req)
    except HTTPException as exc:
        _record_ingest_outcome(req.source_type, modality_label, f"http_{exc.status_code}")
        raise

    try:
        if mode_normalized == "multimodal":
            prepared = _prepare_multimodal_ingest(req)
        elif mode_normalized in {"", "text"}:
            docs = _load_source_documents(req)
            if not docs:
                _record_ingest_metrics(True)
                modality_label = "text"
                _record_ingest_outcome(req.source_type, modality_label, "empty")
                return {"status": "ok", "message": "Aucun document chargé."}
            prepared = _prepare_chunks_for_chroma(req, docs)
        else:
            raise HTTPException(status_code=400, detail="Mode d'ingestion non supporté")
        modality_label = prepared.modality or "unknown"
    except HTTPException as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, f"http_{exc.status_code}")
        raise
    except Exception as exc:
        _record_ingest_metrics(False)
        _record_ingest_outcome(req.source_type, modality_label, "error")
        raise HTTPException(status_code=500, detail=f"Erreur de chargement: {exc}") from exc

    if not prepared.ids:
        _record_ingest_metrics(True)
        _record_ingest_outcome(req.source_type, modality_label, "empty")
        return {"status": "ok", "message": "Aucun contenu éligible à l'ingestion."}

    target_col = resolve_collection_name(section=req.metadata_hints.get("section"), collection=req.metadata_hints.get("collection"))
    return _index_batch(prepared, req.source_type, modality_label, collection_name=target_col)


@app.post("/ingest/urls")
def ingest_urls(
    req: UrlBatchRequest,
    request: Request,
    mode: str = Query(default="text"),
) -> dict[str, Any]:
    """Ingestion par lot d'URLs — vérifie les doublons avant ingestion."""
    _require_api_token_configured()
    _enforce_security(request, req)

    if not req.urls:
        raise HTTPException(status_code=400, detail="La liste d'URLs est vide.")

    results: list[dict[str, Any]] = []
    total_added = 0
    total_skipped = 0

    for url in req.urls:
        url = url.strip()
        if not url:
            continue
        try:
            # Vérifier doublon par content hash
            sub_req = IngestRequest(
                sourceType="url", sourceUrl=url, metadata=req.metadata_hints
            )
            docs = load_from_url(url)
            if not docs:
                results.append({"url": url, "status": "empty", "added": 0, "skipped": 0})
                continue

            prepared = _prepare_chunks_for_chroma(sub_req, docs)
            if not prepared.ids:
                results.append({"url": url, "status": "empty", "added": 0, "skipped": 0})
                continue

            target_col = resolve_collection_name(section=req.metadata_hints.get("section"), collection=req.metadata_hints.get("collection"))
            result = _index_batch(prepared, "url", prepared.modality or "text", collection_name=target_col)
            added = result.get("added", 0)
            skipped = result.get("skipped", 0)
            total_added += added
            total_skipped += skipped
            results.append({"url": url, "status": "ok", "added": added, "skipped": skipped})
        except HTTPException as exc:
            results.append({"url": url, "status": "error", "detail": exc.detail})
        except Exception as exc:
            results.append({"url": url, "status": "error", "detail": str(exc)})

    return {
        "status": "ok",
        "total_urls": len(req.urls),
        "total_added": total_added,
        "total_skipped": total_skipped,
        "results": results,
    }


@app.post("/ingest/upload-files")
async def ingest_upload_files(
    request: Request,
    files: list[UploadFile] = UPLOAD_FILES_PARAM,
    metadata: str = Query(default="{}"),
    mode: str = Query(default="text"),
) -> dict[str, Any]:
    """Upload et ingestion de plusieurs fichiers simultanément avec déduplication."""
    import json as _json
    import uuid as _uuid

    _require_api_token_configured()
    _enforce_security(request, None)

    try:
        hints = _json.loads(metadata)
        if not isinstance(hints, dict):
            hints = {}
    except Exception:
        hints = {}

    upload_dir = Path(os.getenv("ADMIN_UPLOAD_DIR", "/data/uploads"))
    upload_dir.mkdir(parents=True, exist_ok=True)

    results: list[dict[str, Any]] = []
    total_added = 0
    total_skipped = 0
    total_errors = 0

    for file in files:
        fname = file.filename or "upload.bin"
        safe_name = f"{_uuid.uuid4().hex}-{os.path.basename(fname)}"
        dest = upload_dir / safe_name

        try:
            # Sauvegarder le fichier
            content = await file.read()
            dest.write_bytes(content)

            # Calculer le hash pour déduplication
            file_hash = hashlib.sha256(content).hexdigest()

            # Vérifier doublon : même hash déjà présent dans ChromaDB
            target_col = resolve_collection_name(section=hints.get("section"), collection=hints.get("collection"))
            client = get_chroma_client()
            collection = client.get_or_create_collection(
                name=target_col, metadata={"hnsw:space": "cosine"}
            )
            existing = collection.get(where={"file_hash": file_hash}) or {}
            existing_ids = existing.get("ids", []) if existing else []
            if existing_ids:
                skipped = len(existing_ids)
                total_skipped += skipped
                results.append({
                    "filename": fname,
                    "status": "duplicate",
                    "skipped": skipped,
                    "detail": "Fichier déjà ingéré (même empreinte de fichier).",
                })
                try:
                    dest.unlink()
                except OSError:
                    pass
                continue

            # Déterminer le type de source
            ext = dest.suffix.lower()
            source_type_map: dict[str, str] = {
                ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
                ".md": "markdown", ".markdown": "markdown",
                ".txt": "markdown", ".csv": "markdown",
                ".html": "url", ".htm": "url",
                ".jpg": "image", ".jpeg": "image", ".png": "image",
                ".gif": "image", ".bmp": "image", ".webp": "image",
                ".mp3": "audio", ".wav": "audio", ".m4a": "audio", ".ogg": "audio",
                ".flac": "audio", ".aac": "audio",
                ".mp4": "video", ".avi": "video", ".mkv": "video", ".webm": "video", ".mov": "video",
            }
            detected_type = source_type_map.get(ext, "markdown")
            _validate_upload_mode(detected_type, mode)

            # Déterminer le mode d'ingestion
            use_mode = mode

            sub_req = IngestRequest(
                sourceType=cast(Any, detected_type),
                sourceUrl=str(dest),
                metadata={**hints, "original_filename": fname, "file_hash": file_hash},
            )

            if use_mode == "multimodal" and MULTIMODAL_ENABLED:
                prepared = _prepare_multimodal_ingest(sub_req)
            else:
                docs = _load_source_documents(sub_req)
                if not docs:
                    results.append({"filename": fname, "status": "empty", "added": 0, "skipped": 0})
                    continue
                prepared = _prepare_chunks_for_chroma(sub_req, docs)

            if not prepared.ids:
                results.append({"filename": fname, "status": "empty", "added": 0, "skipped": 0})
                continue

            target_col = resolve_collection_name(section=hints.get("section"), collection=hints.get("collection"))
            result = _index_batch(prepared, detected_type, prepared.modality or "text", collection_name=target_col)
            added = result.get("added", 0)
            skipped = result.get("skipped", 0)
            total_added += added
            total_skipped += skipped
            results.append({
                "filename": fname,
                "status": "ok",
                "added": added,
                "skipped": skipped,
                "source_type": detected_type,
            })
        except HTTPException as exc:
            total_errors += 1
            results.append({"filename": fname, "status": "error", "detail": exc.detail})
        except Exception as exc:
            total_errors += 1
            logger.exception("Upload ingestion failed for %s", fname)
            results.append({"filename": fname, "status": "error", "detail": str(exc)})

    return {
        "status": "ok",
        "total_files": len(files),
        "total_added": total_added,
        "total_skipped": total_skipped,
        "total_errors": total_errors,
        "results": results,
    }


@app.post("/ingest/check-duplicates")
def check_duplicates(
    req: DeduplicationCheckRequest,
    request: Request,
) -> dict[str, Any]:
    """Vérifie si des sources ont déjà été ingérées pour éviter les doublons."""
    _require_api_token_configured()
    _enforce_security(request, req)

    target_col = resolve_collection_name(section=req.section, collection=req.collection)
    client = get_chroma_client()
    collection = client.get_or_create_collection(
        name=target_col, metadata={"hnsw:space": "cosine"}
    )

    results: list[dict[str, Any]] = []
    for source in req.sources:
        source = source.strip()
        if not source:
            continue
        # Vérifier par source path dans les métadonnées
        try:
            existing = collection.get(where={"source": source}, limit=1)
            already_ingested = bool(existing and existing.get("ids"))
        except Exception:
            already_ingested = False
        results.append({
            "source": source,
            "already_ingested": already_ingested,
        })

    return {
        "sources_checked": len(results),
        "results": results,
    }


@app.get("/health")
def health_check():
    return {"status": "healthy"}


@app.get("/collections")
def list_collections(request: Request) -> dict[str, Any]:
    """List all ChromaDB collections with document counts and metadata."""
    _require_api_token_configured()
    _enforce_security(request, None)
    try:
        client = get_chroma_client()
        collections_raw = client.list_collections()
        result = []
        for col in collections_raw:
            name = col.name if hasattr(col, 'name') else str(col)
            try:
                c = client.get_collection(name)
                count = c.count()
            except Exception:
                count = 0
            result.append({"name": name, "count": count})
        return {"collections": result, "total": len(result)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Erreur listing collections: {exc}") from exc


@app.get("/stats/{collection_name}")
def collection_stats(collection_name: str, request: Request) -> dict[str, Any]:
    """Get statistics for a specific collection."""
    _require_api_token_configured()
    _enforce_security(request, None)
    target_collection = resolve_collection_name(collection=collection_name, allow_quarantine=True)
    try:
        client = get_chroma_client()
        collection = client.get_or_create_collection(
            name=target_collection, metadata={"hnsw:space": "cosine"}
        )
        count = collection.count()
        # Sample metadata across the full collection to extract unique field values.
        # peek() only returns the first N rows (by insertion order) which may lack metadata.
        # Strategy: fetch all IDs then sample up to 2000 evenly spread across the collection.
        unique_matieres: set[str] = set()
        unique_niveaux: set[str] = set()
        unique_groupes: set[str] = set()
        unique_types: set[str] = set()
        if count > 0:
            try:
                # Stratégie multi-passes : on essaie d'abord un filtre sur source_type (chunks enrichis),
                # puis un peek large en fallback. Ceci évite de charger tous les IDs (timeout sur >100k chunks).
                _meta_sources: list[dict] = []
                # Essayer plusieurs source_type connus pour trouver des chunks avec métadonnées riches.
                # ChromaDB get() ne supporte pas $in — on enchaîne des requêtes exactes.
                for _where in [
                    {"source_type": "gdrive_folder"},
                    {"source_type": "upload"},
                    {"source_type": "url"},
                    {"source_type": "pdf"},
                ]:
                    try:
                        _r = collection.get(where=_where, limit=2000, include=["metadatas"])
                        _found = _r.get("metadatas") or []
                        _meta_sources.extend(_found)
                        if len(_meta_sources) >= 2000:
                            break
                    except Exception:
                        continue
                if not _meta_sources:
                    # Fallback peek
                    _r2 = collection.peek(limit=500)
                    _meta_sources = _r2.get("metadatas") or []
                for m in _meta_sources:
                    if not m:
                        continue
                    if m.get("matiere"):
                        unique_matieres.add(m["matiere"])
                    if m.get("niveau"):
                        unique_niveaux.add(m["niveau"])
                    if m.get("groupe"):
                        unique_groupes.add(m["groupe"])
                    if m.get("type_ressource"):
                        unique_types.add(m["type_ressource"])
            except Exception:
                pass
        
        return {
            "collection": target_collection,
            "doc_count": count,
            "embed_model": EMBED_MODEL,
            "matieres": sorted(unique_matieres),
            "niveaux": sorted(unique_niveaux),
            "groupes": sorted(unique_groupes),
            "types_ressource": sorted(unique_types),
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Erreur stats collection: {exc}") from exc


@app.get("/metrics")
def metrics(request: Request) -> Response:
    _ = request
    if not ingest_metrics.METRICS_ENABLED:
        raise HTTPException(status_code=404, detail="Metrics disabled")
    body = ingest_metrics.generate_latest(METRIC_REGISTRY)
    return Response(body, media_type=CONTENT_TYPE_LATEST)


@app.post("/search")
def search_kb(payload: SearchRequest, request: Request) -> dict[str, Any]:
    # AuthN/AuthZ identical to ingestion
    _require_api_token_configured()
    _enforce_security(request, payload)

    # Prepare chroma collection
    try:
        client = get_chroma_client()
        target_col, effective_filters, maths_fallback_applied = _resolve_search_target(client, payload)
        collection = client.get_or_create_collection(name=target_col, metadata={"hnsw:space": "cosine"})
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail=f"Chroma client error: {exc}") from exc

    # Compute query embedding using the same provider as indexing
    try:
        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        q_vec = emb.embed_query(payload.q)
    except ValueError as exc:
        message = str(exc)
        if "HTTP code: 404" in message:
            logger.warning("Ollama embeddings endpoint returned 404 for model '%s' (search)", EMBED_MODEL)
            raise HTTPException(
                status_code=503,
                detail=(
                    f"Embedding model '{EMBED_MODEL}' is not available on the Ollama backend. "
                    "Pull the model or adjust EMBED_MODEL before retrying."
                ),
            ) from exc
        logger.exception("Embedding provider raised ValueError during search")
        raise HTTPException(status_code=500, detail=f"Embedding error: {message}") from exc
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.exception("Unexpected failure while requesting embeddings (search)")
        raise HTTPException(status_code=500, detail=f"Embedding error: {exc}") from exc

    # Build metadata filters from payload.filters
    # ChromaDB requires $and operator when multiple conditions are present
    where: dict[str, Any] = {}
    if effective_filters:
        conditions = []
        for fk, fv in effective_filters.items():
            if fv is not None and fv != "" and fv != "Tous":
                conditions.append({str(fk): fv})
        if len(conditions) == 1:
            where = conditions[0]
        elif len(conditions) > 1:
            where = {"$and": conditions}

    # Query by embedding
    try:
        n_results = max(1, min(int(payload.k), 50))
        query_kwargs: dict[str, Any] = {"query_embeddings": [q_vec], "n_results": n_results}
        if where:
            query_kwargs["where"] = where
        results = collection.query(**query_kwargs)
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail=f"Chroma query error: {exc}") from exc

    documents = results.get("documents", [[]])[0] if results.get("documents") else []
    metadatas = results.get("metadatas", [[]])[0] if results.get("metadatas") else []
    ids = results.get("ids", [[]])[0] if results.get("ids") else []
    distances = results.get("distances", [[]])[0] if results.get("distances") else []

    hits: list[dict[str, Any]] = []
    for idx, doc_id in enumerate(ids):
        distance = distances[idx] if distances and idx < len(distances) else None
        if (
            payload.score_threshold is not None
            and distance is not None
            and float(distance) > payload.score_threshold
        ):
            continue
        item: dict[str, Any] = {"id": doc_id, "metadata": metadatas[idx] if idx < len(metadatas) else {}}
        if payload.include_documents and idx < len(documents):
            item["document"] = documents[idx]
        if distance is not None:
            item["score"] = distance
        hits.append(item)

    _record_ingest_outcome("search", "text", "success")  # reuse metric surface for visibility
    return {
        "query": payload.q,
        "collection": target_col,
        "k": n_results,
        "returned": len(hits),
        "filters_applied": where,
        "score_threshold": payload.score_threshold,
        "maths_premiere_fallback": maths_fallback_applied,
        "hits": hits,
    }


class RagQueryFilters(BaseModel):
    domain: Optional[str] = None
    document_id: Optional[str] = None
    tags: Optional[list[str]] = None
    metadata: Optional[dict[str, Any]] = None


class RagQuery(BaseModel):
    query: str
    filters: Optional[RagQueryFilters] = None
    top_k: int = Field(default=6, ge=1, le=50)
    collection: str = Field(default=COLLECTION_NAME)


@app.post("/rag/query")
def rag_query(payload: RagQuery, request: Request) -> dict[str, Any]:
    # AuthN/AuthZ identical to ingestion
    _require_api_token_configured()
    _enforce_security(request, payload)

    # Prepare chroma collection
    try:
        client = get_chroma_client()
        target_collection = resolve_collection_name(
            collection=payload.collection,
            allow_quarantine=False,
        )
        collection = client.get_or_create_collection(
            name=target_collection,
            metadata={"hnsw:space": "cosine"},
        )
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail=f"Chroma client error: {exc}") from exc

    # Compute query embedding using the same provider as indexing
    try:
        emb = TimedOllamaEmbeddings(model=EMBED_MODEL, base_url=OLLAMA_URL, request_timeout=OLLAMA_REQUEST_TIMEOUT)
        q_vec = emb.embed_query(payload.query)
    except ValueError as exc:
        message = str(exc)
        if "HTTP code: 404" in message:
            logger.warning("Ollama embeddings endpoint returned 404 for model '%s' (rag/query)", EMBED_MODEL)
            raise HTTPException(
                status_code=503,
                detail=(
                    f"Embedding model '{EMBED_MODEL}' is not available on the Ollama backend. "
                    "Pull the model or adjust EMBED_MODEL before retrying."
                ),
            ) from exc
        logger.exception("Embedding provider raised ValueError during rag/query")
        raise HTTPException(status_code=500, detail=f"Embedding error: {message}") from exc
    except Exception as exc:  # pragma: no cover - defensive logging
        logger.exception("Unexpected failure while requesting embeddings (rag/query)")
        raise HTTPException(status_code=500, detail=f"Embedding error: {exc}") from exc

    # Build metadata filters (where)
    where: dict[str, Any] = {}
    if payload.filters:
        f = payload.filters
        if f.domain:
            where["domain"] = f.domain
        if f.document_id:
            where["document_id"] = f.document_id
        if f.tags:
            where["tags"] = {"$in": f.tags}
        if f.metadata:
            for k, v in (f.metadata or {}).items():
                if v is None or v == "":
                    continue
                where[str(k)] = v

    # Query by embedding with optional filters
    try:
        n_results = max(1, min(int(payload.top_k), 50))
        query_kwargs: dict[str, Any] = {"query_embeddings": [q_vec], "n_results": n_results}
        if where:
            query_kwargs["where"] = where
        results = collection.query(**query_kwargs)
    except Exception as exc:  # pragma: no cover - defensive
        raise HTTPException(status_code=500, detail=f"Chroma query error: {exc}") from exc

    documents = results.get("documents", [[]])[0] if results.get("documents") else []
    metadatas = results.get("metadatas", [[]])[0] if results.get("metadatas") else []
    ids = results.get("ids", [[]])[0] if results.get("ids") else []
    distances = results.get("distances", [[]])[0] if results.get("distances") else []

    hits: list[dict[str, Any]] = []
    for idx, doc_id in enumerate(ids):
        item: dict[str, Any] = {"id": doc_id, "metadata": metadatas[idx] if idx < len(metadatas) else {}}
        if idx < len(documents):
            item["document"] = documents[idx]
        if distances and idx < len(distances) and distances[idx] is not None:
            item["score"] = distances[idx]
        hits.append(item)

    return {
        "query": payload.query,
        "collection": target_collection,
        "k": n_results,
        "filters": where,
        "hits": hits,
    }
